"""Extract readable text from PDF, DOCX, and TXT documents."""

from __future__ import annotations

import io
import re
from pathlib import Path
from typing import BinaryIO, Union

import fitz
from docx import Document


DocumentSource = Union[str, Path, bytes, bytearray, BinaryIO]
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


class DocumentParserError(Exception):
    """Base exception for document parsing problems."""


class UnsupportedFormatError(DocumentParserError):
    """Raised when a document format is not supported."""


class EmptyDocumentError(DocumentParserError):
    """Raised when a document contains no readable text."""


class ExtractionError(DocumentParserError):
    """Raised when a supported document cannot be read or parsed."""


def normalize_text(text: str) -> str:
    """Collapse repeated whitespace while keeping paragraph breaks readable."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    lines = [re.sub(r"[\t \f\v]+", " ", line).strip() for line in text.split("\n")]

    normalized_lines: list[str] = []
    previous_line_was_blank = False
    for line in lines:
        if line:
            normalized_lines.append(line)
            previous_line_was_blank = False
        elif normalized_lines and not previous_line_was_blank:
            normalized_lines.append("")
            previous_line_was_blank = True

    return "\n".join(normalized_lines).strip()


def extract_pdf_text(source: DocumentSource) -> str:
    """Extract text from a PDF document using PyMuPDF."""
    data = _read_bytes(source)
    try:
        with fitz.open(stream=data, filetype="pdf") as document:
            return "\n\n".join(page.get_text() for page in document)
    except Exception as exc:
        raise ExtractionError("Could not extract text from the PDF document.") from exc


def extract_docx_text(source: DocumentSource) -> str:
    """Extract paragraphs and table cells from a DOCX document."""
    data = _read_bytes(source)
    try:
        document = Document(io.BytesIO(data))
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text for cell in row.cells))
        return "\n\n".join(parts)
    except Exception as exc:
        raise ExtractionError("Could not extract text from the DOCX document.") from exc


def extract_txt_text(source: DocumentSource) -> str:
    """Extract UTF-8 text from a TXT document."""
    data = _read_bytes(source)
    try:
        return data.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise ExtractionError("Could not decode the TXT document as UTF-8.") from exc


def extract_text(source: DocumentSource, filename: str | None = None) -> str:
    """Extract and normalize text from a supported document.

    ``source`` may be a filesystem path, bytes, a byte stream, or a Streamlit
    ``UploadedFile``. Raw bytes and unnamed streams require ``filename`` so the
    document format can be identified.
    """
    extension = _get_extension(source, filename)
    extractors = {
        ".pdf": extract_pdf_text,
        ".docx": extract_docx_text,
        ".txt": extract_txt_text,
    }

    if extension not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise UnsupportedFormatError(
            f"Unsupported document format '{extension or 'unknown'}'. "
            f"Supported formats: {supported}."
        )

    try:
        text = normalize_text(extractors[extension](source))
    except DocumentParserError:
        raise
    except Exception as exc:
        raise ExtractionError("The document could not be read.") from exc

    if not text:
        raise EmptyDocumentError("The document contains no readable text.")
    return text


def _get_extension(source: DocumentSource, filename: str | None) -> str:
    """Find a lowercase extension from an explicit or source filename."""
    name = filename
    if name is None and isinstance(source, (str, Path)):
        name = str(source)
    if name is None:
        source_name = getattr(source, "name", None)
        if isinstance(source_name, str):
            name = source_name
    return Path(name).suffix.lower() if name else ""


def _read_bytes(source: DocumentSource) -> bytes:
    """Read document bytes without changing a seekable stream's position."""
    if isinstance(source, (str, Path)):
        try:
            return Path(source).read_bytes()
        except OSError as exc:
            raise ExtractionError(f"Could not read document: {source}") from exc

    if isinstance(source, (bytes, bytearray)):
        return bytes(source)

    if not hasattr(source, "read"):
        raise ExtractionError("Document source must be a path, bytes, or file-like object.")

    original_position = None
    try:
        if hasattr(source, "tell"):
            original_position = source.tell()
        if hasattr(source, "seek"):
            source.seek(0)
        data = source.read()
        if isinstance(data, str):
            data = data.encode("utf-8")
        if not isinstance(data, (bytes, bytearray)):
            raise TypeError("File-like object did not return bytes.")
        return bytes(data)
    except (OSError, TypeError, ValueError) as exc:
        raise ExtractionError("Could not read the document stream.") from exc
    finally:
        if original_position is not None and hasattr(source, "seek"):
            try:
                source.seek(original_position)
            except (OSError, ValueError):
                pass
