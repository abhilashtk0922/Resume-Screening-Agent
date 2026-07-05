"""Tests for PDF, DOCX, and TXT document parsing."""

from io import BytesIO

import fitz
import pytest
from docx import Document

from src.parsers.document_parser import (
    EmptyDocumentError,
    ExtractionError,
    UnsupportedFormatError,
    extract_text,
    normalize_text,
)


class UploadedFileLike(BytesIO):
    """Small stand-in for Streamlit's UploadedFile used in unit tests."""

    def __init__(self, content: bytes, name: str):
        super().__init__(content)
        self.name = name


def test_extracts_txt_from_path(tmp_path):
    document = tmp_path / "resume.txt"
    document.write_text("Jane Doe\n\nPython developer", encoding="utf-8")

    assert extract_text(document) == "Jane Doe\n\nPython developer"


def test_supports_uploaded_file_objects_and_preserves_position():
    uploaded_file = UploadedFileLike(b"Candidate\n\nSkills: Python", "resume.TXT")
    uploaded_file.seek(4)

    assert extract_text(uploaded_file) == "Candidate\n\nSkills: Python"
    assert uploaded_file.tell() == 4


def test_normalizes_excessive_whitespace_and_blank_lines():
    text = "  Jane   Doe  \n\n\n\tPython    Developer\t\n\n\n"

    assert normalize_text(text) == "Jane Doe\n\nPython Developer"


def test_rejects_unsupported_format():
    with pytest.raises(UnsupportedFormatError, match="Unsupported document format"):
        extract_text(b"some content", filename="resume.rtf")


@pytest.mark.parametrize("content", [b"", b"  \n\t\n  "])
def test_rejects_empty_txt_documents(content):
    with pytest.raises(EmptyDocumentError, match="no readable text"):
        extract_text(content, filename="resume.txt")


def test_extracts_pdf_text():
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "PDF Candidate - Python")
    pdf_bytes = document.tobytes()
    document.close()

    assert extract_text(pdf_bytes, filename="resume.pdf") == "PDF Candidate - Python"


def test_extracts_docx_paragraphs_and_tables():
    document = Document()
    document.add_paragraph("DOCX Candidate")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Skill"
    table.cell(0, 1).text = "Python"
    stream = BytesIO()
    document.save(stream)

    assert extract_text(stream.getvalue(), filename="resume.docx") == (
        "DOCX Candidate\n\nSkill | Python"
    )


@pytest.mark.parametrize(
    ("content", "filename"),
    [(b"not a pdf", "resume.pdf"), (b"not a docx", "resume.docx")],
)
def test_corrupted_documents_raise_clear_errors(content, filename):
    with pytest.raises(ExtractionError, match="Could not extract text"):
        extract_text(content, filename=filename)
